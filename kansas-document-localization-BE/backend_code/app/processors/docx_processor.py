import os
from docx import Document

class DOCXProcessor:

    def extract_text(self, input_path: str) -> list:
        segments = []

        document = Document(input_path)

        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()

            if text:
                segments.append({
                    "segment_id": f"docx_p{index}",
                    "page": None,
                    "text": text,
                    "bbox": None,
                    "document_type": "docx"
                })

        return segments

    def rebuild_document(self, input_path: str, translated_segments: list, output_dir: str) -> str:
        document = Document(input_path)

        translation_map = {
            item["segment_id"]: item["translated_text"]
            for item in translated_segments
        }

        for index, paragraph in enumerate(document.paragraphs):
            segment_id = f"docx_p{index}"

            if segment_id in translation_map:
                paragraph.text = translation_map[segment_id]

        output_path = os.path.join(output_dir, "translated_output.docx")
        document.save(output_path)

        return output_path